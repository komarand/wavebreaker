from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from kaggle_researcher.schemas import RetrievedDocument


VALID_NAMING_STRATEGIES = {"timestamp", "increment"}


def generate_report(
    competition_name: str,
    roadmap_text: str,
    sources: list[RetrievedDocument],
    output_path: str | Path,
    *,
    overwrite: bool = False,
    naming_strategy: str = "timestamp",
) -> Path:
    path = resolve_report_path(
        output_path=output_path,
        overwrite=overwrite,
        naming_strategy=naming_strategy,
    )
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading(competition_name, level=0)
    _add_markdown_like_text(document, roadmap_text)

    document.add_heading("Sources", level=1)
    if sources:
        for source in sources:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(source.title).bold = True
            paragraph.add_run(f" [{source.source}]")
            if source.url is not None:
                paragraph.add_run(f" {source.url}")
            quality_score = source.metadata.get("quality_score")
            final_score = source.metadata.get("final_score")
            specificity = source.metadata.get("specificity")
            evidence_type = source.metadata.get("evidence_type")
            if quality_score is not None and final_score is not None:
                paragraph.add_run(
                    " "
                    f"[{specificity or 'unknown'} | {evidence_type or 'unknown'} | "
                    f"final_score={float(final_score):.4f} | rrf={source.rrf_score:.4f} | "
                    f"quality={float(quality_score):.2f}]"
                )
            else:
                paragraph.add_run(f" rrf_score={source.rrf_score:.4f}")
    else:
        document.add_paragraph("No retrieved sources.")

    try:
        document.save(path)
    except PermissionError as exc:
        if overwrite:
            raise RuntimeError(
                f"Cannot overwrite report at {path}. "
                "The file may be open in Word or locked by Windows. "
                "Close it or run without --overwrite-report."
            ) from exc

        fallback_path = make_timestamped_path(path)
        try:
            document.save(fallback_path)
        except PermissionError as fallback_exc:
            raise RuntimeError(
                f"Cannot save report at {path} or fallback path {fallback_path}. "
                "The files may be open in Word or locked by Windows."
            ) from fallback_exc
        path = fallback_path

    return path


def make_timestamped_path(path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    candidate = path.with_name(f"{path.stem}_{timestamp}{path.suffix}")
    if not candidate.exists() and candidate != path:
        return candidate

    counter = 2
    while True:
        incremented = path.with_name(f"{path.stem}_{timestamp}_{counter:03d}{path.suffix}")
        if not incremented.exists() and incremented != path:
            return incremented
        counter += 1


def make_incremented_path(path: Path) -> Path:
    counter = 2
    while True:
        candidate = path.with_name(f"{path.stem}_v{counter:03d}{path.suffix}")
        if not candidate.exists():
            return candidate
        counter += 1


def resolve_report_path(
    output_path: str | Path,
    overwrite: bool = False,
    naming_strategy: str = "timestamp",
) -> Path:
    path = Path(output_path)
    if naming_strategy not in VALID_NAMING_STRATEGIES:
        raise ValueError(
            "naming_strategy must be one of: "
            f"{', '.join(sorted(VALID_NAMING_STRATEGIES))}"
        )

    if overwrite or not path.exists():
        return path

    if naming_strategy == "timestamp":
        return make_timestamped_path(path)
    return make_incremented_path(path)


def _add_markdown_like_text(document: Document, text: str) -> None:
    in_code_block = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if not stripped:
            continue

        if in_code_block:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)
