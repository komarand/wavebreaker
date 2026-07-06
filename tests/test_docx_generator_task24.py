from __future__ import annotations

from pathlib import Path

from docx import Document

from kaggle_researcher.report.docx_generator import generate_report
from kaggle_researcher.schemas import RetrievedDocument


def _source() -> RetrievedDocument:
    return RetrievedDocument(
        id="doc-1",
        competition_id="comp-1",
        source="kaggle",
        title="High scoring notebook",
        url="https://example.com/notebook",
        content="retrieved content",
        score=0.91,
        rrf_score=0.1234,
    )


def test_generate_report_parses_markdown_like_text_and_sources(tmp_path: Path) -> None:
    output_path = tmp_path / "nested" / "report.docx"

    actual_path = generate_report(
        competition_name="comp-1",
        roadmap_text="\n".join(
            [
                "# Roadmap",
                "Intro paragraph.",
                "1. Executive summary",
                "## Experiments",
                "- Build baseline",
                "* Try calibration",
                "```",
                "print('baseline')",
                "```",
                "2) Чего не делать",
                "Avoid public LB chasing.",
            ]
        ),
        sources=[_source()],
        output_path=output_path,
    )

    assert actual_path == output_path
    assert output_path.exists()

    document = Document(output_path)
    styles_by_text = {paragraph.text: paragraph.style.name for paragraph in document.paragraphs}
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    assert styles_by_text["Roadmap"] == "Heading 1"
    assert styles_by_text["Executive summary"] == "Heading 2"
    assert styles_by_text["Experiments"] == "Heading 2"
    assert styles_by_text["Чего не делать"] == "Heading 2"
    assert styles_by_text["Build baseline"] == "List Bullet"
    assert styles_by_text["Try calibration"] == "List Bullet"
    assert "print('baseline')" in paragraph_texts
    assert "Sources" in paragraph_texts
    assert any(
        "High scoring notebook [kaggle] https://example.com/notebook rrf_score=0.1234" in text
        for text in paragraph_texts
    )
