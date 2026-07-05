from __future__ import annotations

from pathlib import Path

from docx import Document

from kaggle_researcher.report import docx_generator
from kaggle_researcher.report.docx_generator import generate_report, resolve_report_path
from kaggle_researcher.schemas import RetrievedDocument


def make_source() -> RetrievedDocument:
    return RetrievedDocument(
        id="doc-1",
        competition_id="comp-1",
        source="kaggle",
        title="High scoring notebook",
        url="https://example.com/notebook",
        content="retrieved content",
        score=0.9,
        rrf_score=0.1234,
    )


def test_generate_report_writes_openable_docx_with_headings_bullets_and_sources(
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "reports" / "comp-1.docx"

    result_path = generate_report(
        competition_name="comp-1",
        roadmap_text="\n".join(
            [
                "# Roadmap",
                "Overview paragraph.",
                "## Experiments",
                "- Build baseline",
                "* Try calibration",
            ]
        ),
        sources=[make_source()],
        output_path=output_path,
    )

    assert result_path == output_path
    assert output_path.exists()
    document = Document(output_path)
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
    styles_by_text = {paragraph.text: paragraph.style.name for paragraph in document.paragraphs}

    assert "comp-1" in paragraph_texts
    assert "Roadmap" in paragraph_texts
    assert "Experiments" in paragraph_texts
    assert "Build baseline" in paragraph_texts
    assert "Try calibration" in paragraph_texts
    assert "Sources" in paragraph_texts
    assert any("High scoring notebook [kaggle]" in text for text in paragraph_texts)
    assert styles_by_text["Build baseline"] == "List Bullet"


def test_generate_report_includes_empty_sources_message(tmp_path: Path) -> None:
    output_path = tmp_path / "empty.docx"

    generate_report("comp-1", "Plain report text.", [], output_path)

    paragraph_texts = [paragraph.text for paragraph in Document(output_path).paragraphs]
    assert "No retrieved sources." in paragraph_texts


def test_resolve_report_path_returns_original_when_missing(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"

    assert resolve_report_path(output_path) == output_path


def test_resolve_report_path_returns_timestamped_path_when_existing(
    monkeypatch,
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "report.docx"
    output_path.write_text("existing", encoding="utf-8")

    class FixedDatetime:
        @classmethod
        def now(cls):
            return cls()

        def strftime(self, fmt: str) -> str:
            return "20260705_143012"

    monkeypatch.setattr(docx_generator, "datetime", FixedDatetime)

    assert resolve_report_path(output_path, naming_strategy="timestamp") == (
        tmp_path / "report_20260705_143012.docx"
    )


def test_resolve_report_path_returns_incremented_path_when_existing(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    output_path.write_text("existing", encoding="utf-8")
    (tmp_path / "report_v002.docx").write_text("existing", encoding="utf-8")

    assert resolve_report_path(output_path, naming_strategy="increment") == (
        tmp_path / "report_v003.docx"
    )


def test_resolve_report_path_overwrite_returns_original_even_when_existing(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"
    output_path.write_text("existing", encoding="utf-8")

    assert resolve_report_path(output_path, overwrite=True) == output_path


class FakeParagraph:
    def add_run(self, text: str = ""):
        return type("FakeRun", (), {"bold": False, "font": type("FakeFont", (), {"name": ""})()})()


class PermissionOnceDocument:
    save_calls: list[Path] = []

    def add_heading(self, *args, **kwargs) -> None:
        return None

    def add_paragraph(self, *args, **kwargs) -> FakeParagraph:
        return FakeParagraph()

    def save(self, path: str | Path) -> None:
        saved_path = Path(path)
        self.save_calls.append(saved_path)
        if len(self.save_calls) == 1:
            raise PermissionError("locked")


class AlwaysLockedDocument(PermissionOnceDocument):
    def save(self, path: str | Path) -> None:
        self.save_calls.append(Path(path))
        raise PermissionError("locked")


def test_generate_report_retries_timestamped_fallback_on_permission_error(
    monkeypatch,
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "report.docx"
    PermissionOnceDocument.save_calls = []

    class FixedDatetime:
        @classmethod
        def now(cls):
            return cls()

        def strftime(self, fmt: str) -> str:
            return "20260705_143012"

    monkeypatch.setattr(docx_generator, "datetime", FixedDatetime)
    monkeypatch.setattr(docx_generator, "Document", PermissionOnceDocument)

    actual_path = generate_report("comp-1", "text", [], output_path)

    assert actual_path == tmp_path / "report_20260705_143012.docx"
    assert PermissionOnceDocument.save_calls == [output_path, actual_path]


def test_generate_report_overwrite_permission_error_is_clear(
    monkeypatch,
    tmp_path: Path,
) -> None:
    output_path = tmp_path / "report.docx"
    AlwaysLockedDocument.save_calls = []
    monkeypatch.setattr(docx_generator, "Document", AlwaysLockedDocument)

    try:
        generate_report("comp-1", "text", [], output_path, overwrite=True)
    except RuntimeError as exc:
        assert "file may be open in Word or locked by Windows" in str(exc)
    else:
        raise AssertionError("Expected RuntimeError")
